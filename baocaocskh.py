import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

def load_app_data(file):
    # Dòng header đúng là dòng số 3 (index 2)
    df = pd.read_excel(file, sheet_name='Sheet1', header=2)
    df.columns = df.columns.str.strip()
    df = df[df['Điện lực'].notna()]
    # Tách dòng tổng hợp "Công ty"
    df_total = df[df['Điện lực'].astype(str).str.strip().str.lower().str.contains('công ty')]
    df_main = df[~df['Điện lực'].astype(str).str.strip().str.lower().str.contains('công ty')]
    # Chuẩn hóa kiểu dữ liệu
    for col in ['STT', 'Số lượng KH quản lý', 'Số lượng đã thực hiện App']:
        df_main[col] = df_main[col].astype(int)
    df_main['Tỷ lệ thực hiện qua App'] = df_main['Tỷ lệ thực hiện qua App'].astype(float)
    if not df_total.empty:
        for col in ['STT', 'Số lượng KH quản lý', 'Số lượng đã thực hiện App']:
            df_total[col] = df_total[col].astype(int)
        df_total['Tỷ lệ thực hiện qua App'] = df_total['Tỷ lệ thực hiện qua App'].astype(float)
    return df_main, df_total

def load_time_data(file):
    # Đọc đúng dòng header (dòng số 4, index 3)
    df = pd.read_excel(file, sheet_name='Sheet1', header=3)
    df.columns = df.columns.str.strip()
    df = df[df['Điện lực'].notna()]
    df_total = df[df['Điện lực'].astype(str).str.strip().str.lower().str.contains('công ty')]
    df_main = df[~df['Điện lực'].astype(str).str.strip().str.lower().str.contains('công ty')]
    for col in ['STT', 'Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn']:
        df_main[col] = df_main[col].astype(int)
    # Xử lý tỷ lệ trễ hạn dạng "0.02%" hoặc số thực
    df_main['Tỷ lệ trễ hạn'] = df_main['Tỷ lệ trễ hạn'].astype(str).str.replace('%','').astype(float)
    if not df_total.empty:
        for col in ['Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn']:
            df_total[col] = df_total[col].astype(int)
        df_total['Tỷ lệ trễ hạn'] = df_total['Tỷ lệ trễ hạn'].astype(str).str.replace('%','').astype(float)
    return df_main, df_total

def plot_bar(df, x_col, y_col, title, ylabel, percent=False, color='skyblue'):
    fig, ax = plt.subplots(figsize=(7,4))
    bars = ax.bar(df[x_col], df[y_col], color=color, alpha=0.85)
    ax.set_ylabel(ylabel)
    ax.set_title(title)
    plt.xticks(rotation=30, ha='right')
    if percent:
        labels = [f"{v*100:.2f}%" for v in df[y_col]]
    else:
        labels = [f"{v:.2f}%" for v in df[y_col]]
    ax.bar_label(bars, labels=labels, fontsize=11)
    plt.tight_layout()
    return fig

def add_table(doc, df, columns, percent_col=None):
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light List Accent 1'
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = row[col]
            if percent_col and col == percent_col:
                row_cells[i].text = f"{float(val)*100:.2f}"
            elif isinstance(val, float) and not col == percent_col:
                row_cells[i].text = f"{val:.0f}" if val.is_integer() else f"{val:.2f}"
            else:
                row_cells[i].text = str(val)

def export_word_app(app_df, app_total, app_info, app_top3, app_bottom3, app_fig_top, app_fig_bot):
    doc = Document()
    doc.add_heading('BÁO CÁO ĐÁNH GIÁ KẾT QUẢ THỰC HIỆN QUA APP CSKH', 0)
    doc.add_heading('1. Bảng dữ liệu tổng hợp', level=1)
    columns = ['STT', 'Điện lực', 'Số lượng KH quản lý', 'Số lượng đã thực hiện App', 'Tỷ lệ thực hiện qua App']
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light List Accent 1'
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in app_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số lượng KH quản lý']))
        row_cells[3].text = str(int(row['Số lượng đã thực hiện App']))
        row_cells[4].text = f"{row['Tỷ lệ thực hiện qua App']*100:.2f}%"
    # Thêm dòng Công ty nếu có
    if not app_total.empty:
        row = app_total.iloc[0]
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số lượng KH quản lý']))
        row_cells[3].text = str(int(row['Số lượng đã thực hiện App']))
        row_cells[4].text = f"{row['Tỷ lệ thực hiện qua App']*100:.2f}%"
    doc.add_heading('2. Nhận xét tổng quan', level=1)
    doc.add_paragraph(f"Tổng số khách hàng quản lý: {app_info['total_kh']:,}")
    doc.add_paragraph(f"Tổng số khách hàng đã thực hiện App: {app_info['total_app']:,}")
    doc.add_paragraph(f"Tỷ lệ thực hiện qua App trung bình: {app_info['avg_rate']*100:.2f}%")
    doc.add_paragraph(
        f"Tỷ lệ thực hiện qua App trung bình đạt {app_info['avg_rate']*100:.2f}%. "
        "Một số điện lực đạt tỷ lệ cao, cho thấy hiệu quả truyền thông và hỗ trợ khách hàng tốt. "
        "Các đơn vị có tỷ lệ thấp cần tăng cường truyền thông, hỗ trợ kỹ thuật và khuyến khích khách hàng sử dụng App."
    )
    doc.add_heading('3. Top 3 điện lực tỷ lệ thực hiện qua App cao nhất:', level=1)
    add_table(doc, app_top3, ['STT', 'Điện lực', 'Tỷ lệ thực hiện qua App', 'Số lượng đã thực hiện App', 'Số lượng KH quản lý'], percent_col='Tỷ lệ thực hiện qua App')
    doc.add_paragraph(
        f"Điện lực {app_top3.iloc[0]['Điện lực']} dẫn đầu với tỷ lệ {app_top3.iloc[0]['Tỷ lệ thực hiện qua App']*100:.2f}%. "
        f"Các điện lực top 3 đều có tỷ lệ trên {app_top3['Tỷ lệ thực hiện qua App'].min()*100:.2f}%."
    )
    doc.add_paragraph("Biểu đồ Top 3:")
    buf1 = BytesIO()
    app_fig_top.savefig(buf1, format='png')
    buf1.seek(0)
    doc.add_picture(buf1, width=Inches(5))

    doc.add_heading('4. Top 3 điện lực tỷ lệ thực hiện qua App thấp nhất:', level=1)
    add_table(doc, app_bottom3, ['STT', 'Điện lực', 'Tỷ lệ thực hiện qua App', 'Số lượng đã thực hiện App', 'Số lượng KH quản lý'], percent_col='Tỷ lệ thực hiện qua App')
    doc.add_paragraph(
        f"Các điện lực nhóm cuối như {', '.join(app_bottom3['Điện lực'])} có tỷ lệ thấp hơn đáng kể, cần tập trung cải thiện."
    )
    doc.add_paragraph("Biểu đồ Bottom 3:")
    buf2 = BytesIO()
    app_fig_bot.savefig(buf2, format='png')
    buf2.seek(0)
    doc.add_picture(buf2, width=Inches(5))
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def export_word_time(time_df, time_total, time_info, time_top3, time_bottom3, time_fig_top, time_fig_bot):
    doc = Document()
    doc.add_heading('BÁO CÁO ĐÁNH GIÁ YÊU CẦU GIẢI QUYẾT ĐÚNG THỜI GIAN CAM KẾT', 0)
    doc.add_heading('1. Bảng dữ liệu tổng hợp', level=1)
    columns = ['STT', 'Điện lực', 'Số yêu cầu chuyển xử lý', 'Số lượng phiếu giải quyết trễ hạn', 'Tỷ lệ trễ hạn']
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Light List Accent 1'
    for i, col in enumerate(columns):
        table.rows[0].cells[i].text = col
    for _, row in time_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(int(row['STT']))
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số yêu cầu chuyển xử lý']))
        row_cells[3].text = str(int(row['Số lượng phiếu giải quyết trễ hạn']))
        row_cells[4].text = f"{row['Tỷ lệ trễ hạn']:.2f}%"
    if not time_total.empty:
        row = time_total.iloc[0]
        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = str(row['Điện lực'])
        row_cells[2].text = str(int(row['Số yêu cầu chuyển xử lý']))
        row_cells[3].text = str(int(row['Số lượng phiếu giải quyết trễ hạn']))
        row_cells[4].text = f"{row['Tỷ lệ trễ hạn']:.2f}%"
    doc.add_heading('2. Nhận xét tổng quan', level=1)
    doc.add_paragraph(f"Tổng số yêu cầu chuyển xử lý: {time_info['total_req']:,}")
    doc.add_paragraph(f"Tổng số phiếu giải quyết trễ hạn: {time_info['total_late']:,}")
    doc.add_paragraph(f"Tỷ lệ trễ hạn trung bình: {time_info['avg_late']:.2f}%")
    doc.add_paragraph(
        f"Tỷ lệ trễ hạn trung bình là {time_info['avg_late']:.2f}%. "
        "Hầu hết các điện lực duy trì tỷ lệ trễ hạn thấp, nhưng vẫn còn một số đơn vị tỷ lệ cao hơn mức trung bình."
    )
    doc.add_heading('3. Top 3 điện lực tỷ lệ trễ hạn cao nhất:', level=1)
    add_table(doc, time_top3, ['STT', 'Điện lực', 'Tỷ lệ trễ hạn', 'Số lượng phiếu giải quyết trễ hạn', 'Số yêu cầu chuyển xử lý'], percent_col='Tỷ lệ trễ hạn')
    doc.add_paragraph(
        f"Điện lực {time_top3.iloc[0]['Điện lực']} có tỷ lệ trễ hạn cao nhất ({time_top3.iloc[0]['Tỷ lệ trễ hạn']:.2f}%). "
        f"Các đơn vị top 3 đều có tỷ lệ trên {time_top3['Tỷ lệ trễ hạn'].min():.2f}%."
    )
    doc.add_paragraph("Biểu đồ Top 3:")
    buf3 = BytesIO()
    time_fig_top.savefig(buf3, format='png')
    buf3.seek(0)
    doc.add_picture(buf3, width=Inches(5))

    doc.add_heading('4. Top 3 điện lực tỷ lệ trễ hạn thấp nhất:', level=1)
    add_table(doc, time_bottom3, ['STT', 'Điện lực', 'Tỷ lệ trễ hạn', 'Số lượng phiếu giải quyết trễ hạn', 'Số yêu cầu chuyển xử lý'], percent_col='Tỷ lệ trễ hạn')
    doc.add_paragraph(
        f"Các đơn vị như {', '.join(time_bottom3['Điện lực'])} duy trì tỷ lệ rất thấp, là điểm sáng cần nhân rộng."
    )
    doc.add_paragraph("Biểu đồ Bottom 3:")
    buf4 = BytesIO()
    time_fig_bot.savefig(buf4, format='png')
    buf4.seek(0)
    doc.add_picture(buf4, width=Inches(5))
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

st.title("Báo cáo đánh giá kết quả thực hiện & Xuất file Word")

tab1, tab2 = st.tabs(["Đánh giá qua App CSKH", "Đánh giá đúng thời gian cam kết"])

with tab1:
    st.header("Đánh giá kết quả thực hiện qua App CSKH")
    file_app = st.file_uploader("Chọn file báo cáo App CSKH", type=["xlsx"], key="app")
    if file_app:
        df_app, df_app_total = load_app_data(file_app)
        st.subheader("Bảng dữ liệu tổng hợp")
        df_app_show = pd.concat([df_app, df_app_total], ignore_index=True)
        df_app_show['Tỷ lệ thực hiện qua App (%)'] = df_app_show['Tỷ lệ thực hiện qua App']*100
        st.dataframe(df_app_show.drop(columns=['Tỷ lệ thực hiện qua App']), hide_index=True)
        total_kh = df_app_show['Số lượng KH quản lý'].sum()
        total_app = df_app_show['Số lượng đã thực hiện App'].sum()
        avg_rate = df_app['Tỷ lệ thực hiện qua App'].mean()
        st.markdown("**Nhận xét tổng quan:**")
        st.info(
            f"Tổng số khách hàng quản lý: {total_kh:,}\n\n"
            f"Tổng số khách hàng đã thực hiện App: {total_app:,}\n\n"
            f"Tỷ lệ thực hiện qua App trung bình: {avg_rate*100:.2f}%\n\n"
            f"Tỷ lệ thực hiện qua App trung bình đạt {avg_rate*100:.2f}%. "
            "Một số điện lực đạt tỷ lệ cao, cho thấy hiệu quả truyền thông và hỗ trợ khách hàng tốt. "
            "Các đơn vị có tỷ lệ thấp cần tăng cường truyền thông, hỗ trợ kỹ thuật và khuyến khích khách hàng sử dụng App."
        )
        # Top 3
        top3 = df_app.nlargest(3, 'Tỷ lệ thực hiện qua App')
        st.subheader("Top 3 điện lực tỷ lệ thực hiện qua App cao nhất")
        st.dataframe(
            top3[['STT', 'Điện lực', 'Tỷ lệ thực hiện qua App', 'Số lượng đã thực hiện App', 'Số lượng KH quản lý']]
            .assign(**{'Tỷ lệ thực hiện qua App (%)': top3['Tỷ lệ thực hiện qua App']*100})
            .drop('Tỷ lệ thực hiện qua App',axis=1)
            .reset_index(drop=True), hide_index=True)
        st.markdown(
            f"Điện lực {top3.iloc[0]['Điện lực']} dẫn đầu với tỷ lệ {top3.iloc[0]['Tỷ lệ thực hiện qua App']*100:.2f}%. "
            f"Các điện lực top 3 đều có tỷ lệ trên {top3['Tỷ lệ thực hiện qua App'].min()*100:.2f}%."
        )
        fig_top3 = plot_bar(top3, 'Điện lực', 'Tỷ lệ thực hiện qua App', "Top 3 tỷ lệ thực hiện qua App", "Tỷ lệ thực hiện qua App (%)", percent=True, color='royalblue')
        st.pyplot(fig_top3)
        # Bottom 3
        bottom3 = df_app.nsmallest(3, 'Tỷ lệ thực hiện qua App')
        st.subheader("Top 3 điện lực tỷ lệ thực hiện qua App thấp nhất")
        st.dataframe(
            bottom3[['STT', 'Điện lực', 'Tỷ lệ thực hiện qua App', 'Số lượng đã thực hiện App', 'Số lượng KH quản lý']]
            .assign(**{'Tỷ lệ thực hiện qua App (%)': bottom3['Tỷ lệ thực hiện qua App']*100})
            .drop('Tỷ lệ thực hiện qua App',axis=1)
            .reset_index(drop=True), hide_index=True)
        st.markdown(
            f"Các điện lực nhóm cuối như {', '.join(bottom3['Điện lực'])} có tỷ lệ thấp hơn đáng kể, cần tập trung cải thiện."
        )
        fig_bot3 = plot_bar(bottom3, 'Điện lực', 'Tỷ lệ thực hiện qua App', "Bottom 3 tỷ lệ thực hiện qua App", "Tỷ lệ thực hiện qua App (%)", percent=True, color='orange')
        st.pyplot(fig_bot3)
        # Nút xuất Word cho tab 1
        if st.button("Tải báo cáo Word tab này", key="word_app"):
            word_file = export_word_app(df_app, df_app_total, {'total_kh': total_kh, 'total_app': total_app, 'avg_rate': avg_rate}, top3, bottom3, fig_top3, fig_bot3)
            st.download_button(label="Tải báo cáo Word", data=word_file, file_name="Bao_cao_AppCSKH.docx")

with tab2:
    st.header("Đánh giá kết quả giải quyết yêu cầu đúng thời gian cam kết")
    file_time = st.file_uploader("Chọn file báo cáo trễ hạn", type=["xlsx"], key="time")
    if file_time:
        df_time, df_time_total = load_time_data(file_time)
        st.subheader("Bảng dữ liệu tổng hợp")
        df_time_show = pd.concat([df_time, df_time_total], ignore_index=True)
        df_time_show['Tỷ lệ trễ hạn (%)'] = df_time_show['Tỷ lệ trễ hạn']
        st.dataframe(df_time_show.drop(columns=['Tỷ lệ trễ hạn']), hide_index=True)
        total_req = df_time_show['Số yêu cầu chuyển xử lý'].sum()
        total_late = df_time_show['Số lượng phiếu giải quyết trễ hạn'].sum()
        avg_late = df_time['Tỷ lệ trễ hạn'].mean()
        st.markdown("**Nhận xét tổng quan:**")
        st.info(
            f"Tổng số yêu cầu chuyển xử lý: {total_req:,}\n\n"
            f"Tổng số phiếu giải quyết trễ hạn: {total_late:,}\n\n"
            f"Tỷ lệ trễ hạn trung bình: {avg_late:.2f}%\n\n"
            f"Tỷ lệ trễ hạn trung bình là {avg_late:.2f}%. "
            "Hầu hết các điện lực duy trì tỷ lệ trễ hạn thấp, nhưng vẫn còn một số đơn vị tỷ lệ cao hơn mức trung bình."
        )
        # Top 3
        top3 = df_time.nlargest(3, 'Tỷ lệ trễ hạn')
        st.subheader("Top 3 điện lực tỷ lệ trễ hạn cao nhất")
        st.dataframe(
            top3[['STT', 'Điện lực', 'Tỷ lệ trễ hạn', 'Số lượng phiếu giải quyết trễ hạn', 'Số yêu cầu chuyển xử lý']]
            .rename(columns={'Tỷ lệ trễ hạn':'Tỷ lệ trễ hạn (%)'})
            .reset_index(drop=True), hide_index=True)
        st.markdown(
            f"Điện lực {top3.iloc[0]['Điện lực']} có tỷ lệ trễ hạn cao nhất ({top3.iloc[0]['Tỷ lệ trễ hạn']:.2f}%). "
            f"Các đơn vị top 3 đều có tỷ lệ trên {top3['Tỷ lệ trễ hạn'].min():.2f}%."
        )
        fig_top3 = plot_bar(top3, 'Điện lực', 'Tỷ lệ trễ hạn', "Top 3 tỷ lệ trễ hạn", "Tỷ lệ trễ hạn (%)", percent=False, color='crimson')
        st.pyplot(fig_top3)
        # Bottom 3
        bottom3 = df_time.nsmallest(3, 'Tỷ lệ trễ hạn')
        st.subheader("Top 3 điện lực tỷ lệ trễ hạn thấp nhất")
        st.dataframe(
            bottom3[['STT', 'Điện lực', 'Tỷ lệ trễ hạn', 'Số lượng phiếu giải quyết trễ hạn', 'Số yêu cầu chuyển xử lý']]
            .rename(columns={'Tỷ lệ trễ hạn':'Tỷ lệ trễ hạn (%)'})
            .reset_index(drop=True), hide_index=True)
        st.markdown(
            f"Các đơn vị như {', '.join(bottom3['Điện lực'])} duy trì tỷ lệ rất thấp, là điểm sáng cần nhân rộng."
        )
        fig_bot3 = plot_bar(bottom3, 'Điện lực', 'Tỷ lệ trễ hạn', "Bottom 3 tỷ lệ trễ hạn", "Tỷ lệ trễ hạn (%)", percent=False, color='goldenrod')
        st.pyplot(fig_bot3)
        # Nút xuất Word cho tab 2
        if st.button("Tải báo cáo Word tab này", key="word_time"):
            word_file = export_word_time(df_time, df_time_total, {'total_req': total_req, 'total_late': total_late, 'avg_late': avg_late}, top3, bottom3, fig_top3, fig_bot3)
            st.download_button(label="Tải báo cáo Word", data=word_file, file_name="Bao_cao_TreHan.docx")
